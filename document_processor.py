import hashlib
import magic
import io
import shutil
import json
import time
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from pypdf import PdfReader
from docx import Document
from PIL import Image
from bs4 import BeautifulSoup
import markdown
import chardet
import base64
import pandas as pd
from models import ProcessedFile, DocumentChunk, ImageMetadata
from langchain_text_splitters import RecursiveCharacterTextSplitter
from security import RateLimiter, validate_content
from langchain_core.messages import HumanMessage
from llm_provider import get_llm

class DocumentProcessor:
    """Handles processing of different document types"""
    
    # Initialize rate limiter as a class variable
    _rate_limiter = RateLimiter()
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        
        self.mime = magic.Magic(mime=True)
        
        # Initialize vision model
        self.vision_model = None
        vision_model = "llama3.2-vision:latest"
        
        try:
            print(f"Attempting to initialize vision model {vision_model}...")
            self.vision_model = get_llm(vision_model)
            print(f"Successfully initialized vision model {vision_model}")
        except Exception as e:
            print(f"Failed to initialize {vision_model}: {e}")
            print("WARNING: Failed to initialize vision model. Image processing will be unavailable.")
    
    def process_file(self, file_path: Path, original_filename: str = None) -> Optional[ProcessedFile]:
        """Process a document or image file and return chunks with metadata.
        
        Handles multiple file types:
        - Documents (PDF, DOCX, TXT): Extracts and chunks text
        - Images: Uses vision model to extract text and create chunks
        
        For images, the process includes:
        - Extracting text using Llava vision model
        - Creating unified chunks for vector storage
        - Preserving image metadata (dimensions, format)
        
        Args:
            file_path: Path to the file to process
            original_filename: Optional original name of the file
            
        Returns:
            ProcessedFile with chunks and metadata, or None if processing fails
        """
        try:
            # Check rate limit
            self._rate_limiter.check_rate_limit('process_file')
            
            # Get original filename if not provided
            if not original_filename:
                original_filename = file_path.name
            
            # Process file
            print(f"Processing file: {file_path}")
            file_type = self._get_file_type(file_path)
            print(f"Detected file type: {file_type}")
            content, image_metadata = self._extract_content(file_path, file_type)
            print(f"Extracted content length: {len(content) if content else 0}, Has image metadata: {bool(image_metadata)}")
            
            # Validate content
            if not content or not content.strip():
                raise ValueError("No valid content extracted from file")
                
            # Create chunks based on file type
            if image_metadata:
                # For images, use the chunks from metadata
                chunks = [DocumentChunk(
                    text=chunk,
                    metadata={
                        "source": original_filename or str(file_path),
                        "page": "1",  # Images are single page
                        "chunk_type": "image_text",
                        "chunk_index": i,
                        "width": image_metadata.width,
                        "height": image_metadata.height,
                        "format": image_metadata.format
                    }
                ) for i, chunk in enumerate(image_metadata.chunks)]
            else:
                # For documents, split content into chunks
                chunks = self._split_content(content, file_path, file_type, original_filename)
            
            checksum = self._calculate_checksum(file_path)
            
            return ProcessedFile(
                filename=original_filename,
                file_type=file_type,
                checksum=checksum,
                chunks=chunks,
                metadata={"content": content} if bool(image_metadata) else {},
                is_image=bool(image_metadata),
                image_metadata=image_metadata
            )
        except ValueError as ve:
            print(f"Error processing file: {str(ve)}")
            return None
        except Exception as e:
            print(f"Unexpected error processing file: {str(e)}")
            return None
    
    def _get_file_type(self, file_path: Path) -> str:
        """Determine file type from extension"""
        ext = file_path.suffix.lower()[1:] if file_path.suffix else ''
        print(f"Detecting file type for extension: '{ext}'")
        
        # Map extensions to types
        if ext in ['txt', 'md', 'html']:
            return ext
        elif ext == 'pdf':
            return 'pdf'
        elif ext == 'docx':
            return 'docx'
        elif ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
            return 'image'
        elif ext in ['xlsx', 'xls', 'csv']:
            return 'spreadsheet'
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _calculate_checksum(self, file_path: Path) -> str:
        """Calculate SHA-256 checksum of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def _extract_content(self, file_path: Path, file_type: str) -> Tuple[str, Optional[ImageMetadata]]:
        """Extract text content from file based on type.
        
        Args:
            file_path: Path to the file to process
            file_type: Type of file (pdf, docx, txt, etc)
            
        Returns:
            Tuple containing:
            - Extracted text content
            - Optional ImageMetadata for images
            
        Raises:
            ValueError: If file is empty or cannot be processed
        """
        # Check if file exists and is not empty
        if not file_path.exists():
            raise ValueError(f"File does not exist: {file_path}")
        if file_path.stat().st_size == 0:
            raise ValueError(f"File is empty: {file_path}")
            
        # Check if file is an image
        mime_type = self.mime.from_file(str(file_path))
        print(f"Detected MIME type: {mime_type}")
        if mime_type.startswith('image/'):
            print("File identified as image, using vision model processing")
            text, metadata = self._extract_image(file_path)
            if not metadata:
                raise ValueError(f"Failed to analyze image: {file_path.name}")
            if not text.strip():
                raise ValueError(f"No text extracted from image: {file_path.name}")
            return text, metadata
            
        # Handle other file types
        content = None
        if file_type == 'pdf':
            content = self._extract_pdf(file_path)
        elif file_type == 'docx':
            content = self._extract_docx(file_path)
        elif file_type == 'spreadsheet':
            content = self._extract_spreadsheet(file_path)
        elif file_type in ['txt', 'md', 'html']:
            content = self._extract_text(file_path, file_type)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
            
        # Validate extracted content
        if not content or not content.strip():
            raise ValueError(f"No valid content extracted from {file_type} file: {file_path.name}")
            
        return content, None
    
    def _split_content(self, content: str, file_path: Path, file_type: str, original_filename: str = None) -> List[DocumentChunk]:
        """Split content into chunks with metadata.
        
        Args:
            content: Text content to split
            file_path: Path to the source file
            file_type: Type of file (pdf, docx, txt, etc)
            original_filename: Optional original name of the file
            
        Returns:
            List of DocumentChunks with metadata
        """
        # Split content into chunks
        texts = self.text_splitter.split_text(content)
        
        # Create chunks with metadata
        chunks = []
        for i, text in enumerate(texts):
            if not text.strip():  # Skip empty chunks
                continue
                
            chunks.append(DocumentChunk(
                text=text,
                metadata={
                    "source": original_filename or str(file_path),
                    "chunk_type": file_type,
                    "chunk_index": i,
                    "page": str(i + 1)  # Each chunk gets its own page number
                }
            ))
        
        return chunks
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF.
        
        Extracts text while preserving:
        - Page breaks
        - Basic formatting
        - Tables and lists (as text)
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text with preserved formatting
            
        Raises:
            ValueError: If file is empty or cannot be processed
        """
        try:
            # Validate file
            if not file_path.exists():
                raise ValueError(f"PDF file does not exist: {file_path}")
            if file_path.stat().st_size == 0:
                raise ValueError(f"PDF file is empty: {file_path}")
                
            # Open PDF
            try:
                reader = PdfReader(file_path)
            except Exception as e:
                raise ValueError(f"Failed to open PDF file: {str(e)}")
                
            if not reader.pages:
                raise ValueError("PDF file has no pages")
                
            content = []
            
            # Process each page
            for i, page in enumerate(reader.pages):
                try:
                    # Extract text with layout preservation
                    text = page.extract_text()
                    if not text:  # Skip empty pages
                        continue
                        
                    # Clean and normalize text
                    text = text.strip()
                    text = '\n'.join(line.strip() for line in text.splitlines() if line.strip())
                    
                    if text:  # Only add if we have content
                        content.append(f"[Page {i+1}]\n{text}")
                        
                except Exception as e:
                    print(f"Warning: Failed to process page {i+1}: {str(e)}")
                    continue
            
            # Join and validate final content
            text = '\n\n'.join(content)
            if not text.strip():
                raise ValueError("No valid content extracted from PDF file")
                
            return text
            
        except ValueError as ve:
            raise ve  # Re-raise validation errors
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file.
        
        Extracts text while preserving:
        - Paragraphs with proper spacing
        - Tables (converts to text)
        - Lists and numbering
        - Headers and sections
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text with preserved formatting
            
        Raises:
            ValueError: If file is empty or cannot be processed
        """
        try:
            # Validate file
            if not file_path.exists():
                raise ValueError(f"DOCX file does not exist: {file_path}")
            if file_path.stat().st_size == 0:
                raise ValueError(f"DOCX file is empty: {file_path}")
                
            # Open document
            try:
                doc = Document(file_path)
            except Exception as e:
                raise ValueError(f"Failed to open DOCX file: {str(e)}")
                
            content = []
            
            # Extract headers
            for section in doc.sections:
                try:
                    header = section.header
                    if header.text and header.text.strip():
                        content.append(header.text.strip())
                except Exception as e:
                    print(f"Warning: Failed to extract header: {str(e)}")
                    continue
            
            # Extract main content with style preservation
            for paragraph in doc.paragraphs:
                try:
                    text = paragraph.text.strip()
                    if not text:  # Skip empty paragraphs
                        continue
                        
                    # Handle list items and special styles
                    style = paragraph.style.name.lower()
                    if 'list' in style or 'bullet' in style or (paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None):
                        content.append('  • ' + text)
                    elif 'heading' in style:
                        content.append('\n' + text + '\n')
                    else:
                        content.append(text)
                except Exception as e:
                    print(f"Warning: Failed to process paragraph: {str(e)}")
                    continue
            
            # Extract tables
            for table in doc.tables:
                try:
                    # Add table header
                    content.append('\nTable:')
                    for row in table.rows:
                        try:
                            cells = [cell.text.strip() for cell in row.cells if cell]
                            if any(cells):  # Skip empty rows
                                row_text = ' | '.join(cell for cell in cells if cell)
                                content.append(row_text)
                        except Exception as e:
                            print(f"Warning: Failed to process table row: {str(e)}")
                            continue
                    content.append('')  # Add space after table
                except Exception as e:
                    print(f"Warning: Failed to process table: {str(e)}")
                    continue
            
            # Join and validate final content
            text = '\n\n'.join(content)
            if not text.strip():
                raise ValueError("No valid content extracted from DOCX file")
                
            return text
            
        except ValueError as ve:
            raise ve  # Re-raise validation errors
        except Exception as e:
            raise ValueError(f"Failed to process DOCX file: {str(e)}")
    
    def _extract_image(self, file_path: Path) -> Tuple[str, ImageMetadata]:
        """Extract text and metadata from image file using the vision model.
        
        Uses a multi-stage process:
        1. Extract basic image metadata (dimensions, format)
        2. Process image with vision model for text extraction
        3. Create text chunks for vector storage
        4. Store all metadata including extracted text
        
        Args:
            file_path: Path to the image file
            
        Returns:
            Tuple containing:
            - Extracted text as a string
            - ImageMetadata with dimensions, format, and text chunks
        """
        print(f"Starting image extraction for {file_path}")
        try:
            # Validate file
            print("Validating image file...")
            if not file_path.exists():
                raise ValueError(f"Image file does not exist: {file_path}")
            if file_path.stat().st_size == 0:
                raise ValueError(f"Image file is empty: {file_path}")
            print(f"Image file validated successfully. Size: {file_path.stat().st_size} bytes")
                
            # Create temp directory for image processing if it doesn't exist
            print("Setting up temp directory...")
            temp_dir = Path(".temp_images")
            temp_dir.mkdir(exist_ok=True)
            
            # Copy image to temp directory with unique name
            target_path = temp_dir / f"{hashlib.md5(str(file_path).encode()).hexdigest()}_{file_path.name}"
            target_path.parent.mkdir(exist_ok=True)
            print(f"Created temp path: {target_path}")
            
            try:
                shutil.copy2(file_path, target_path)
            except Exception as e:
                raise ValueError(f"Failed to copy image file: {str(e)}")
            
            # Get image metadata and validate format
            try:
                # Open image and keep it open until we're done with metadata
                img = Image.open(target_path)
                print(f"Image format detected: {img.format}")
                # Convert MPO to JPEG if needed
                if img.format == 'MPO':
                    print(f"Converting MPO to JPEG...")
                    # MPO files contain multiple images, we'll use the first one
                    img = img.convert('RGB')
                    # Save to a new path with .jpg extension
                    jpg_path = target_path.with_suffix('.jpg')
                    img.save(jpg_path, format='JPEG', quality=95)
                    print(f"Successfully converted MPO to JPEG at {jpg_path}")
                    # Update target_path to point to the new jpg file
                    target_path = jpg_path
                    # Reopen the converted image to verify
                    img = Image.open(target_path)
                    print(f"Reopened image format: {img.format}")
                # Validate format for other types
                elif img.format not in ['JPEG', 'PNG', 'GIF', 'BMP']:
                    raise ValueError(f"Unsupported image format: {img.format}")
                    
                # Validate dimensions
                if img.width < 10 or img.height < 10:
                    raise ValueError(f"Image too small: {img.width}x{img.height}")
                if img.width > 4096 or img.height > 4096:
                    # Resize large images
                    ratio = min(4096/img.width, 4096/img.height)
                    new_width = int(img.width * ratio)
                    new_height = int(img.height * ratio)
                    img = img.resize((new_width, new_height))
                    
                    # Save resized image
                    img.save(target_path, format='JPEG', quality=95)
                    width, height = new_width, new_height
                else:
                    width, height = img.width, img.height
                    
                # Convert to RGB if needed
                if img.mode != 'RGB':
                    print(f"Converting image from {img.mode} to RGB")
                    img = img.convert('RGB')
                    print(f"Successfully converted to RGB mode")
                
                print(f"Saving image to temp file: {target_path}")
                img.save(target_path, format='JPEG', quality=95)
                print("Image saved successfully")
                    
                # Create metadata
                print("Creating image metadata")
                print(f"Final image mode: {img.mode}")
                metadata = ImageMetadata(
                    width=width,
                    height=height,
                    format='jpeg',  # We always save as JPEG
                    mode=img.mode,  # Should be 'RGB' after conversion
                    has_text=False,
                    detected_objects=[],
                    frame_count=getattr(img, 'n_frames', 1)
                )
                print(f"Metadata created: {metadata}")
                
                # Close the image
                img.close()
            except Exception as e:
                if target_path.exists():
                    target_path.unlink()
                raise ValueError(f"Failed to process image file: {str(e)}")
            
            # Convert image to base64 for vision model
            try:
                print("Converting image to base64...")
                with open(target_path, 'rb') as img_file:
                    img_data = img_file.read()
                    print(f"Read {len(img_data)} bytes from image file")
                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                    print(f"Base64 encoding successful, length: {len(img_base64)}")
            except Exception as e:
                print(f"Error during base64 encoding: {str(e)}")
                if target_path.exists():
                    target_path.unlink()
                raise ValueError(f"Failed to encode image: {str(e)}")
            
            # Process with vision model
            if not self.vision_model:
                print("ERROR: Vision model not initialized")
                if target_path.exists():
                    target_path.unlink()
                raise ValueError("Vision model not initialized")
            
            print(f"Using vision model: {self.vision_model}")
            
            # Prepare request for vision model
            prompt = "Extract and list ALL text from this image exactly as it appears. Include every word, number, and symbol. Preserve line breaks and spacing."
            print("Prepared vision model prompt")
            
            try:
                # Make request with retries
                max_retries = 3
                text = ""
                for attempt in range(max_retries):
                    try:
                        print(f"Attempt {attempt + 1} to process image...")
                        
                        # Create message with image for Ollama vision model
                        message = HumanMessage(
                            content=[
                                {
                                    "type": "text",
                                    "text": prompt
                                },
                                {
                                    "type": "image_url",
                                    "image_url": f"data:image/jpeg;base64,{img_base64}"
                                }
                            ]
                        )
                        
                        # Get response from vision model
                        print("Sending request to vision model...")
                        response = self.vision_model.invoke([message])
                        print(f"Got response from vision model: {response}")
                        
                        if hasattr(response, 'content'):
                            text = response.content.strip()
                        else:
                            text = str(response).strip()
                        
                        if text:
                            print(f"Successfully extracted text: {text[:100]}...")
                            break
                        else:
                            print("No text extracted from response")
                            
                        if attempt < max_retries - 1:
                            print("Waiting before retry...")
                            time.sleep(2)  # Longer wait between retries
                            
                    except Exception as e:
                        if attempt == max_retries - 1:
                            raise e
                        print(f"Warning: Vision model attempt {attempt + 1} failed: {str(e)}")
                        time.sleep(1)  # Wait before retry
                
                # Clean up the image file
                if target_path.exists():
                    target_path.unlink()
                    
                # Update metadata with extracted text
                metadata.detected_text = text
                metadata.has_text = bool(text)
                
                # Create comprehensive text representation
                text_parts = [
                    f"Image File: {file_path.name}",
                    f"Dimensions: {metadata.width}x{metadata.height}",
                    f"Format: {metadata.format}",
                    "\nExtracted Text:",
                    text
                ]
                
                # Join parts and create chunks
                full_text = '\n'.join(text_parts)
                
                # Split into chunks if we have content
                chunks = self.text_splitter.split_text(full_text) if text else []
                metadata.chunks = chunks
                
                return full_text, metadata
                
            except Exception as e:
                if target_path.exists():
                    target_path.unlink()
                print(f"Warning: Vision model failed: {str(e)}")
                
                # Create default metadata if none exists
                if metadata is None:
                    metadata = ImageMetadata(
                        width=0,
                        height=0,
                        format="unknown",
                        mode="unknown",
                        has_text=False,
                        detected_objects=[],
                        frame_count=1,
                        vision_error=str(e)
                    )
                return "", metadata
                
        except ValueError as ve:
            if target_path.exists():
                target_path.unlink()
            raise ve  # Re-raise validation errors
        except Exception as e:
            if target_path.exists():
                target_path.unlink()
            raise ValueError(f"Failed to process image: {str(e)}")
    
    def _extract_spreadsheet(self, file_path: Path) -> str:
        """Extract text from spreadsheet files (Excel, CSV).
        
        Handles:
        - Excel files (XLSX, XLS)
        - CSV files
        - Preserves sheet names, headers, and data structure
        
        Args:
            file_path: Path to the spreadsheet file
            
        Returns:
            Extracted text with preserved structure
            
        Raises:
            ValueError: If file is empty or cannot be processed
        """
        try:
            # Determine file type
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
                sheets = {'Sheet1': df}  # CSV files have one sheet
            else:
                # Excel files can have multiple sheets
                sheets = pd.read_excel(file_path, sheet_name=None)
            
            # Process each sheet
            text_parts = []
            for sheet_name, df in sheets.items():
                if df.empty:
                    continue
                    
                # Add sheet name as header
                text_parts.append(f"Sheet: {sheet_name}\n{'='*40}\n")
                
                # Convert headers and data to string format
                headers = df.columns.tolist()
                text_parts.append(f"Headers: {', '.join(str(h) for h in headers)}\n")
                
                # Process each row
                for idx, row in df.iterrows():
                    row_text = [f"{headers[i]}: {str(val)}" for i, val in enumerate(row)]
                    text_parts.append(f"Row {idx + 1}:\n" + '\n'.join(row_text) + '\n')
                
                text_parts.append('\n' + '-'*40 + '\n')  # Sheet separator
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            raise ValueError(f"Failed to process spreadsheet: {str(e)}")

    def _extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract text from text-based files with proper formatting preservation.
        
        Handles:
        - Text files with various encodings
        - Markdown with formatting
        - HTML with structure
        
        Args:
            file_path: Path to the text file
            file_type: Type of text file (txt, md, html)
            
        Returns:
            Extracted text with preserved formatting
            
        Raises:
            ValueError: If file is empty or cannot be processed
        """
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                if not raw_data.strip():  # Check for empty file
                    raise ValueError("File is empty")
                    
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
                if not content.strip():  # Check for whitespace-only
                    raise ValueError("File contains only whitespace")
            
            # Process based on file type
            if file_type == 'md':
                # Convert markdown to HTML while preserving structure
                html = markdown.markdown(
                    content,
                    extensions=['tables', 'fenced_code', 'nl2br']
                )
                # Convert HTML to text while preserving structure
                soup = BeautifulSoup(html, 'html.parser')
                # Handle lists and code blocks specially
                for elem in soup.find_all(['pre', 'code']):
                    elem.insert_before('\n')
                    elem.insert_after('\n')
                for elem in soup.find_all(['li']):
                    elem.insert_before('  • ')
                text = soup.get_text(separator='\n\n')
            
            elif file_type == 'html':
                # Parse HTML and preserve structure
                soup = BeautifulSoup(content, 'html.parser')
                # Remove script and style elements
                for elem in soup(['script', 'style']):
                    elem.decompose()
                # Handle lists specially
                for elem in soup.find_all(['li']):
                    elem.insert_before('  • ')
                text = soup.get_text(separator='\n\n')
            
            else:  # Plain text
                text = content
            
            # Final cleanup
            text = '\n'.join(line.strip() for line in text.splitlines())
            text = '\n'.join(chunk for chunk in text.split('\n\n') if chunk.strip())
            
            if not text.strip():
                raise ValueError("No valid text content extracted")
                
            return text
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from {file_type} file: {str(e)}")
